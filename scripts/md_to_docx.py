#!/usr/bin/env python3
"""
Convert Markdown to Word (.docx) preserving structure and formatting.
Handles headers, tables, code blocks, lists, bold, links, and horizontal rules.
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_horizontal_line(paragraph):
    """Add a horizontal line after a paragraph."""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)


def parse_table(lines):
    """Parse markdown table lines into rows of cells."""
    rows = []
    for line in lines:
        # Skip separator line (|---|---|)
        if re.match(r'^\s*\|[\s\-:]+\|', line):
            continue
        # Get cells between first and last |
        parts = line.split('|')
        if len(parts) >= 2:
            cells = [c.strip() for c in parts[1:-1]]
            if any(cells):
                rows.append(cells)
    return rows


def add_bold_parts(paragraph, text):
    """Add text to paragraph, handling **bold** as runs."""
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def convert_md_to_docx(md_path: Path, docx_path: Path) -> None:
    """Convert Markdown file to Word document."""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_block_lines = []
    code_lang = ''
    in_table = False
    table_lines = []

    while i < len(lines):
        line = lines[i]
        line_stripped = line.strip()

        # Code block handling
        if line_stripped.startswith('```'):
            if not in_code_block:
                # Flush table if any
                if in_table and table_lines:
                    rows = parse_table(table_lines)
                    if rows:
                        max_cols = max(len(r) for r in rows)
                        table = doc.add_table(rows=len(rows), cols=max_cols)
                        table.style = 'Table Grid'
                        for ri, row_data in enumerate(rows):
                            padded = row_data + [''] * (max_cols - len(row_data))
                            for ci, cell_data in enumerate(padded):
                                if ci < len(table.rows[ri].cells):
                                    table.rows[ri].cells[ci].text = cell_data
                    in_table = False
                    table_lines = []

                in_code_block = True
                code_lang = line_stripped[3:].strip()
                code_block_lines = []
            else:
                # End of code block
                code_text = '\n'.join(code_block_lines)
                p = doc.add_paragraph()
                p.style = 'Normal'
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                in_code_block = False
            i += 1
            continue

        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue

        # Table handling
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                if table_lines:
                    rows = parse_table(table_lines)
                    if rows:
                        max_cols = max(len(r) for r in rows)
                        table = doc.add_table(rows=len(rows), cols=max_cols)
                        table.style = 'Table Grid'
                        for ri, row_data in enumerate(rows):
                            padded = row_data + [''] * (max_cols - len(row_data))
                            for ci, cell_data in enumerate(padded):
                                if ci < len(table.rows[ri].cells):
                                    table.rows[ri].cells[ci].text = cell_data
                    table_lines = []
                in_table = True
            table_lines.append(line)
            i += 1
            continue
        else:
            if in_table and table_lines:
                rows = parse_table(table_lines)
                if rows:
                    max_cols = max(len(r) for r in rows)
                    table = doc.add_table(rows=len(rows), cols=max_cols)
                    table.style = 'Table Grid'
                    for ri, row_data in enumerate(rows):
                        padded = row_data + [''] * (max_cols - len(row_data))
                        for ci, cell_data in enumerate(padded):
                            if ci < len(table.rows[ri].cells):
                                table.rows[ri].cells[ci].text = cell_data
                in_table = False
                table_lines = []

        # Horizontal rule
        if line_stripped in ('---', '***', '___'):
            p = doc.add_paragraph()
            add_horizontal_line(p)
            i += 1
            continue

        # Headers
        if line_stripped.startswith('#'):
            match = re.match(r'^(#{1,4})\s+(.+)$', line_stripped)
            if match:
                level = len(match.group(1))
                text = match.group(2).strip()
                if level == 1:
                    p = doc.add_heading(text, level=0)
                    p.runs[0].font.size = Pt(16)
                else:
                    doc.add_heading(text, level=min(level, 3))
                i += 1
                continue

        # Numbered list
        num_match = re.match(r'^(\d+)\.\s+(.+)$', line_stripped)
        if num_match:
            p = doc.add_paragraph(style='List Number')
            add_bold_parts(p, num_match.group(2))
            i += 1
            continue

        # Bullet list
        if line_stripped.startswith('- ') or line_stripped.startswith('* '):
            text = line_stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            add_bold_parts(p, text)
            i += 1
            continue

        # Empty line
        if not line_stripped:
            doc.add_paragraph()
            i += 1
            continue

        # Regular paragraph (may contain **bold** and [links](url))
        p = doc.add_paragraph()
        # Simple inline formatting
        remaining = line_stripped
        while remaining:
            bold_match = re.search(r'\*\*([^*]+)\*\*', remaining)
            link_match = re.search(r'\[([^\]]+)\]\(([^)]+)\)', remaining)
            
            # Find earliest match
            pos = len(remaining)
            match_type = None
            match_obj = None
            if bold_match and bold_match.start() < pos:
                pos = bold_match.start()
                match_type = 'bold'
                match_obj = bold_match
            if link_match and link_match.start() < pos:
                pos = link_match.start()
                match_type = 'link'
                match_obj = link_match

            if pos > 0:
                p.add_run(remaining[:pos])
                remaining = remaining[pos:]
            elif match_type == 'bold':
                run = p.add_run(match_obj.group(1))
                run.bold = True
                remaining = remaining[match_obj.end():]
            elif match_type == 'link':
                p.add_run(match_obj.group(1) + ' ')
                remaining = remaining[match_obj.end():]
            else:
                p.add_run(remaining)
                remaining = ''

        i += 1

    # Flush any remaining table
    if table_lines:
        rows = parse_table(table_lines)
        if rows:
            max_cols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=max_cols)
            table.style = 'Table Grid'
            for ri, row_data in enumerate(rows):
                padded = row_data + [''] * (max_cols - len(row_data))
                for ci, cell_data in enumerate(padded):
                    if ci < len(table.rows[ri].cells):
                        table.rows[ri].cells[ci].text = cell_data

    doc.save(docx_path)
    print(f"Created: {docx_path}")


if __name__ == '__main__':
    base = Path(__file__).parent.parent
    md_file = base / 'docs' / 'Detailed_Design_Document.md'
    docx_file = base / 'docs' / 'Detailed_Design_Document.docx'
    if len(sys.argv) >= 2:
        md_file = Path(sys.argv[1])
    if len(sys.argv) >= 3:
        docx_file = Path(sys.argv[2])
    else:
        docx_file = md_file.with_suffix('.docx')
    convert_md_to_docx(md_file, docx_file)
